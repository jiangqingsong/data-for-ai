"""数据 Pipeline — 文档加载、清洗、分块、向量化

支持 PDF（文字版/扫描版 OCR）和 Word (.docx) 文档。
"""
import os
import re
from typing import List

import fitz  # pymupdf
import numpy as np
from PIL import Image
from langchain_core.documents import Document

# EasyOCR 全局实例，延迟初始化
_ocr_reader = None


def _get_ocr_reader():
    """延迟初始化 EasyOCR reader（只在首次遇到扫描页时加载模型）"""
    global _ocr_reader
    if _ocr_reader is None:
        import easyocr
        _ocr_reader = easyocr.Reader(["ch_sim", "en"], gpu=True)
    return _ocr_reader


def _is_scanned_page(text: str, threshold: int = 50) -> bool:
    """判断页面是否为扫描版（有效文本不足）"""
    cleaned = text.strip()
    for wm in ["微信公众号：电子课本大全", "电子课本大全"]:
        cleaned = cleaned.replace(wm, "")
    return len(cleaned.strip()) < threshold


def _ocr_page(pixmap, ocr_reader) -> str:
    """对单页渲染图片执行 OCR，返回识别文本"""
    img = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
    img_np = np.array(img)
    results = ocr_reader.readtext(img_np)
    lines = [item[1] for item in results]
    return "\n".join(lines)


def _load_docx(filepath: str) -> str:
    """加载 Word (.docx) 文件，提取纯文本"""
    from docx import Document as DocxDocument
    doc = DocxDocument(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def _is_docx(filename: str) -> bool:
    return filename.lower().endswith(".docx")


def load_documents(doc_dir: str, file_filter: str | None = None) -> List[Document]:
    """加载目录下的 PDF 和 Word 文档

    PDF：逐页提取文本，自动检测扫描页并 OCR
    Word：提取段落文本，作为一个 Document

    Args:
        doc_dir: 文档目录
        file_filter: 文件名过滤关键字，None=全部

    每个 Document 包含:
      - page_content: 文本内容
      - metadata: {"source": 文件名, "page": 页码, "is_scanned": bool}
    """
    all_docs = []
    supported_exts = (".pdf", ".docx")
    doc_files = sorted([
        f for f in os.listdir(doc_dir)
        if f.lower().endswith(supported_exts) and not f.startswith("~")
    ])

    if file_filter:
        doc_files = [f for f in doc_files if file_filter in f]

    if not doc_files:
        print(f"警告: {doc_dir} 目录下没有找到可处理的文档")
        return all_docs

    ocr_reader = None
    scanned_count = 0
    docx_count = 0

    for doc_file in doc_files:
        filepath = os.path.join(doc_dir, doc_file)
        print(f"  加载: {doc_file} ...")

        if _is_docx(doc_file):
            # Word 文档处理
            try:
                text = _load_docx(filepath)
                all_docs.append(Document(
                    page_content=text,
                    metadata={
                        "source": doc_file,
                        "page": 0,
                        "is_scanned": False,
                    },
                ))
                docx_count += 1
                print(f"    OK: Word 文档, {len(text)} 字符")
            except Exception as e:
                print(f"    FAIL: {doc_file}, 错误: {e}")
            continue

        # PDF 文档处理
        try:
            doc = fitz.open(filepath)
            text_pages = 0
            ocr_pages = 0

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()

                if _is_scanned_page(text):
                    if ocr_reader is None:
                        print("    检测到扫描页，加载 OCR 模型...")
                        ocr_reader = _get_ocr_reader()
                    pix = page.get_pixmap(dpi=200)
                    text = _ocr_page(pix, ocr_reader)
                    ocr_pages += 1
                    scanned_count += 1
                else:
                    text_pages += 1

                all_docs.append(Document(
                    page_content=text,
                    metadata={
                        "source": doc_file,
                        "page": page_num,
                        "is_scanned": _is_scanned_page(page.get_text()),
                    },
                ))

            doc.close()
            print(f"    OK: 文字 {text_pages} 页 + OCR {ocr_pages} 页")

        except Exception as e:
            print(f"    FAIL: {doc_file}, 错误: {e}")

    print(f"共加载 {len(doc_files)} 个文档, {len(all_docs)} 页"
          f" (PDF {len(doc_files) - docx_count} + Word {docx_count}, OCR {scanned_count} 页)")
    return all_docs


# 保持向后兼容的别名
load_pdfs = load_documents


WATERMARK_PATTERNS = [
    "微信公众号：电子课本大全",
    "电子课本大全",
    "微信公众号",
]


def clean_documents(docs: List[Document]) -> List[Document]:
    """清洗文档文本：去水印、去除多余空白、规范化中文

    - 移除水印文字
    - 合并连续空格
    - 合并连续换行（最多保留1个）
    """
    for doc in docs:
        text = doc.page_content
        # 移除水印
        for wm in WATERMARK_PATTERNS:
            text = text.replace(wm, "")
        # 合并连续空格
        text = re.sub(r"[ \t]+", " ", text)
        # 合并连续换行（保留最多1个换行）
        text = re.sub(r"\n{3,}", "\n\n", text)
        # 去除首尾空白
        text = text.strip()
        doc.page_content = text

    # 过滤掉清洗后为空的文档
    docs = [doc for doc in docs if len(doc.page_content) > 10]
    return docs


# ============================================================
# Task 3: 文档分块 + 向量化存储
# ============================================================

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma


def split_documents(
    docs: List[Document],
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
) -> List[Document]:
    """将文档按指定大小分块

    Args:
        docs: 原始文档列表
        chunk_size: 每块最大字符数
        chunk_overlap: 块间重叠字符数

    Returns:
        分块后的文档列表，保留原始元数据
    """
    # 中文友好的分隔符：段落 → 换行 → 句号 → 空格
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", ".", " ", ""],
        length_function=len,
    )
    chunks = text_splitter.split_documents(docs)
    print(f"分块完成: {len(docs)} 页面 → {len(chunks)} 个块")
    return chunks


def build_vectorstore(
    docs: List[Document],
    persist_dir: str = "../data/chroma_db",
    embedding_model: str | None = None,
    batch_size: int = 100,
) -> Chroma:
    """将文档向量化并分批存入 Chroma

    火山引擎 Embedding API 限制每次最多 256 条输入，
    因此需要分批处理。

    Args:
        docs: 待向量化的文档列表
        persist_dir: Chroma 持久化目录
        embedding_model: None=自动读取 Config, 或手动指定模型名
        batch_size: 每批处理的文档数（默认100，最大256）

    Returns:
        Chroma vectorstore 实例
    """
    from src.config import Config
    from langchain_openai import OpenAIEmbeddings

    if embedding_model is None:
        embedding_model = Config.EMBEDDING_API_MODEL

    embeddings = OpenAIEmbeddings(
        model=embedding_model,
        openai_api_key=Config.DEEPSEEK_API_KEY,
        openai_api_base=Config.DEEPSEEK_BASE_URL,
        tiktoken_enabled=False,
        check_embedding_ctx_length=False,
    )

    # 分批处理：先创建空 Chroma，再逐批添加
    total = len(docs)
    vectorstore = None

    for i in range(0, total, batch_size):
        batch = docs[i : i + batch_size]
        batch_end = min(i + batch_size, total)
        print(f"  向量化批次 [{i+1}-{batch_end}] / {total}...")

        if vectorstore is None:
            # 第一批：创建 Chroma 并写入
            vectorstore = Chroma.from_documents(
                documents=batch,
                embedding=embeddings,
                persist_directory=persist_dir,
            )
        else:
            # 后续批次：追加
            vectorstore.add_documents(batch)

    print(f"向量化完成: {total} 个块 → Chroma ({persist_dir})")
    return vectorstore


def run_pipeline(
    pdf_dir: str | None = None,
    persist_dir: str | None = None,
    subdir: str | None = None,
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    file_filter: str | None = None,
) -> Chroma:
    """一键执行完整的数据 Pipeline

    PDF/Word → 加载 → 清洗 → 分块 → 向量化 → Chroma

    Args:
        pdf_dir: 文档目录（None=自动从 Config 读取）
        persist_dir: Chroma 持久化目录（None=自动从 Config 读取）
        subdir: PDF_BASE_DIR 下的子目录名，文档和向量库按子目录分存
        file_filter: 文件名过滤关键字，None=全部
    """
    from src.config import Config

    if pdf_dir is None:
        pdf_dir = Config.get_pdf_dir(subdir)
    if persist_dir is None:
        persist_dir = Config.get_chroma_dir(subdir)

    print("=" * 50)
    print("数据 Pipeline 开始执行")
    if subdir:
        print(f"子目录: {subdir}")
    print(f"文档目录: {pdf_dir}")
    print(f"向量库:   {persist_dir}")
    print("=" * 50)

    # 1. 加载
    print("\n[1/4] 加载文档...")
    docs = load_documents(pdf_dir, file_filter=file_filter)
    if not docs:
        raise ValueError(f"未找到可处理的文档: {pdf_dir}")

    # 2. 清洗
    print("\n[2/4] 文本清洗...")
    docs = clean_documents(docs)

    # 3. 分块
    print(f"\n[3/4] 文档分块 (chunk_size={chunk_size}, overlap={chunk_overlap})...")
    chunks = split_documents(docs, chunk_size, chunk_overlap)

    # 4. 向量化
    print(f"\n[4/4] 向量化存储...")
    vectorstore = build_vectorstore(
        chunks, persist_dir
    )

    print("\n" + "=" * 50)
    print("Pipeline 执行完成!")
    print(f"  输入: {len(docs)} 个页面")
    print(f"  输出: {len(chunks)} 个向量块")
    print(f"  存储: {persist_dir}")
    print("=" * 50)
    return vectorstore


if __name__ == "__main__":
    import sys
    import argparse

    parser = argparse.ArgumentParser(description="RAG 数据 Pipeline — PDF/Word → 向量库")
    parser.add_argument(
        "file_filter", nargs="?", default=None,
        help="文件名过滤关键字（如 '数学'），只处理含关键字的文档"
    )
    parser.add_argument(
        "--subdir", default=None,
        help="PDF_BASE_DIR 下的子目录名，PDF 和向量库按子目录分存"
    )
    args = parser.parse_args()

    if args.file_filter:
        print(f"[filter] 过滤关键字: {args.file_filter}")
    if args.subdir:
        print(f"[subdir] 子目录: {args.subdir}")

    run_pipeline(subdir=args.subdir, file_filter=args.file_filter)


# ============================================================
# Task 8: 高级分块策略
# ============================================================

def semantic_split(
    docs: List[Document],
    embedding_model: str | None = None,
    threshold_percentile: float = 90,
) -> List[Document]:
    """语义分块 — 基于 Embedding 相似度骤降点切分

    原理:
      1. 将文档拆成句子
      2. 计算相邻句子的 Embedding 余弦相似度
      3. 相似度低于阈值（分位点）的位置 = 切分点

    Args:
        docs: 原始文档列表
        embedding_model: Embedding 模型名 (None=读 Config)
        threshold_percentile: 相似度阈值分位点，越大切得越细

    Returns:
        语义分块后的文档列表
    """
    from langchain_experimental.text_splitter import SemanticChunker
    from src.config import Config
    from langchain_openai import OpenAIEmbeddings

    if embedding_model is None:
        embedding_model = Config.EMBEDDING_API_MODEL

    embeddings = OpenAIEmbeddings(
        model=embedding_model,
        openai_api_key=Config.DEEPSEEK_API_KEY,
        openai_api_base=Config.DEEPSEEK_BASE_URL,
        tiktoken_enabled=False,
        check_embedding_ctx_length=False,
    )

    splitter = SemanticChunker(
        embeddings=embeddings,
        breakpoint_threshold_type="percentile",
        breakpoint_threshold_amount=threshold_percentile,
    )
    chunks = splitter.split_documents(docs)
    print(f"语义分块完成: {len(docs)} 页面 → {len(chunks)} 个块")
    return chunks


def structural_split(
    docs: List[Document],
) -> List[Document]:
    """按文档结构分块 — 利用标题层级切分

    尝试检测 Markdown 标题结构来切分。对于纯文本 PDF，
    会尝试识别章节标题模式（如"第X章""第X节"）作为切分边界。

    Args:
        docs: 原始文档列表

    Returns:
        结构分块后的文档列表
    """
    import re
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    # 中文教材常见的章节标题模式
    chapter_patterns = [
        r"第[一二三四五六七八九十\d]+章",   # 第X章
        r"第\d+节",                          # 第X节
    ]

    all_chunks = []
    for doc in docs:
        text = doc.page_content
        # 尝试按章节标题切分
        combined_pattern = "(" + "|".join(chapter_patterns) + ")"
        sections = re.split(combined_pattern, text)

        if len(sections) <= 1:
            # 没有匹配到章节结构，整页作为一个块
            if len(text.strip()) > 10:
                all_chunks.append(doc)
            continue

        # 重新组装：标题 + 内容
        i = 0
        while i < len(sections):
            section_text = sections[i]
            if re.match(combined_pattern, section_text) and i + 1 < len(sections):
                # 标题 + 下一段内容
                chunk_text = section_text + sections[i + 1]
                i += 2
            else:
                chunk_text = section_text
                i += 1

            if len(chunk_text.strip()) > 10:
                from langchain_core.documents import Document as Doc
                new_doc = Doc(
                    page_content=chunk_text.strip(),
                    metadata=doc.metadata.copy(),
                )
                all_chunks.append(new_doc)

    # 对于过大的块，用 RecursiveCharacterTextSplitter 二次切分
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,
        chunk_overlap=200,
        separators=["\n\n", "\n", "。", ".", " ", ""],
        length_function=len,
    )
    final_chunks = splitter.split_documents(all_chunks)

    print(f"结构分块完成: {len(docs)} 页面 → {len(final_chunks)} 个块")
    return final_chunks


def parent_document_split(
    docs: List[Document],
    parent_chunk_size: int = 2000,
    child_chunk_size: int = 500,
    child_chunk_overlap: int = 100,
) -> tuple[List[Document], List[Document]]:
    """小2大分块 — 父块用于生成, 子块用于检索

    Args:
        docs: 原始文档列表
        parent_chunk_size: 父块大小（用于生成 context）
        child_chunk_size: 子块大小（用于向量检索）
        child_chunk_overlap: 子块重叠大小

    Returns:
        (parent_docs, child_docs) 元组
    """
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    # 父块切分器（大块）
    parent_splitter = RecursiveCharacterTextSplitter(
        chunk_size=parent_chunk_size,
        chunk_overlap=int(parent_chunk_size * 0.1),
        separators=["\n\n", "\n", "。", ".", " ", ""],
        length_function=len,
    )

    # 子块切分器（小块）
    child_splitter = RecursiveCharacterTextSplitter(
        chunk_size=child_chunk_size,
        chunk_overlap=child_chunk_overlap,
        separators=["\n\n", "\n", "。", ".", " ", ""],
        length_function=len,
    )

    parent_docs = parent_splitter.split_documents(docs)
    child_docs = child_splitter.split_documents(docs)

    # 为子块建立到父块的映射关系
    for i, parent in enumerate(parent_docs):
        parent.metadata["parent_id"] = f"parent_{i}"

    for child in child_docs:
        # 找到包含此子块的父块（通过内容前50字符匹配）
        for i, parent in enumerate(parent_docs):
            if child.page_content[:50] in parent.page_content:
                child.metadata["parent_id"] = f"parent_{i}"
                break

    print(
        f"小2大分块: {len(docs)} 页面"
        f" → {len(parent_docs)} 父块 + {len(child_docs)} 子块"
    )
    return parent_docs, child_docs


def run_chunking_experiment(
    docs: List[Document],
) -> dict:
    """分块策略对比实验 — 一次运行，四种策略，记录统计信息

    Returns:
        {"fixed": {chunks, avg_size}, "semantic": {...},
         "structural": {...}, "parent_child": {...}}
    """
    results = {}

    # 策略1: 固定大小分块（基线）
    print("\n--- 策略1: 固定大小分块 ---")
    fixed_chunks = split_documents(docs, chunk_size=1000, chunk_overlap=200)
    results["fixed"] = {
        "chunks": len(fixed_chunks),
        "avg_size": round(
            sum(len(c.page_content) for c in fixed_chunks)
            / max(len(fixed_chunks), 1)
        ),
    }
    print(f"  → {len(fixed_chunks)} chunks, 平均 {results['fixed']['avg_size']} 字符")

    # 策略2: 语义分块
    print("\n--- 策略2: 语义分块 ---")
    try:
        semantic_chunks = semantic_split(docs)
        results["semantic"] = {
            "chunks": len(semantic_chunks),
            "avg_size": round(
                sum(len(c.page_content) for c in semantic_chunks)
                / max(len(semantic_chunks), 1)
            ),
        }
        print(f"  → {len(semantic_chunks)} chunks, 平均 {results['semantic']['avg_size']} 字符")
    except Exception as e:
        print(f"  语义分块失败: {e}")
        results["semantic"] = {"error": str(e)}

    # 策略3: 结构分块
    print("\n--- 策略3: 结构分块 ---")
    try:
        structural_chunks = structural_split(docs)
        results["structural"] = {
            "chunks": len(structural_chunks),
            "avg_size": round(
                sum(len(c.page_content) for c in structural_chunks)
                / max(len(structural_chunks), 1)
            ),
        }
        print(f"  → {len(structural_chunks)} chunks, 平均 {results['structural']['avg_size']} 字符")
    except Exception as e:
        print(f"  结构分块失败: {e}")
        results["structural"] = {"error": str(e)}

    # 策略4: 小2大分块
    print("\n--- 策略4: 小2大分块 ---")
    parent_docs, child_docs = parent_document_split(docs)
    results["parent_child"] = {
        "parents": len(parent_docs),
        "children": len(child_docs),
        "parent_avg_size": round(
            sum(len(p.page_content) for p in parent_docs)
            / max(len(parent_docs), 1)
        ),
        "child_avg_size": round(
            sum(len(c.page_content) for c in child_docs)
            / max(len(child_docs), 1)
        ),
    }
    print(
        f"  → {len(parent_docs)} 父块 (avg {results['parent_child']['parent_avg_size']} 字符)"
        f" + {len(child_docs)} 子块 (avg {results['parent_child']['child_avg_size']} 字符)"
    )

    # 打印汇总
    print("\n" + "=" * 55)
    print("分块策略对比汇总")
    print("=" * 55)
    print(f"  {'策略':<14} {'块数':<8} {'平均大小':<10}")
    print(f"  {'-'*32}")
    for name, label in [
        ("fixed", "固定大小"),
        ("semantic", "语义分块"),
        ("structural", "结构分块"),
    ]:
        info = results.get(name, {})
        if "error" in info:
            print(f"  {label:<14} {'ERROR':<8} {info['error'][:30]}")
        else:
            print(f"  {label:<14} {info.get('chunks', '?'):<8} {info.get('avg_size', '?'):<10}")
    pc = results.get("parent_child", {})
    print(f"  {'小2大(子)':<14} {pc.get('children', '?'):<8} {pc.get('child_avg_size', '?'):<10}")
    print(f"  {'小2大(父)':<14} {pc.get('parents', '?'):<8} {pc.get('parent_avg_size', '?'):<10}")
    print("=" * 55)

    return results
